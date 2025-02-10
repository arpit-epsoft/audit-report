from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
import time
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import TimeoutException
import undetected_chromedriver as uc
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement  # Ensure this import is included
import pandas as pd
from openai import OpenAI
import requests
import os
from dotenv import load_dotenv

load_dotenv()


def check_robots_txt_and_sitemap(site_url):
    options = uc.ChromeOptions()
    # options.add_argument("--headless=new")
    # options.add_argument("--no-sandbox")  # Recommended for some environments
    # options.add_argument("--disable-dev-shm-usage")
    options.add_argument("user-data-dir=C:/Users/user/AppData/Local/Google/Chrome/User Data")
    browser = uc.Chrome(options=options)

        # Create the full URL for robots.txt
    robots_url = site_url.rstrip('/') + '/robots.txt'
        
    try:
        # Make a GET request to the robots.txt URL
        response = requests.get(robots_url)

        if response.status_code == 200:
            robots_status = f"Yes, robots.txt is present. Status Code: {response.status_code}"
            time.sleep(2)
            # Extract sitemap URLs using regex
            sitemap_urls = re.findall(r"Sitemap:\s*(\S+)", response.text, re.IGNORECASE)
            sitemap_count = len(sitemap_urls)

            if sitemap_count > 0:
                sitemap_status = f"Found {sitemap_count} sitemap(s). URLs: {', '.join(sitemap_urls)}"
            else:
                sitemap_status = "No sitemap URLs found in robots.txt."
            
        else:
            robots_status = f"No, robots.txt is not present. Status Code: {response.status_code}"
            sitemap_status = "N/A"

    except requests.exceptions.RequestException as e:
        robots_status = f"Error: {e}"
        sitemap_status = "N/A"

    finally:
        browser.quit()
        
    return {
        robots_status, sitemap_status
    }

# Function to scrape the Speed Index from PageSpeed Insights
def get_speed_index(company_url):
    # Setup undetected ChromeDriver
    options = uc.ChromeOptions()
    # options.add_argument("--headless=new")
    # options.add_argument("--no-sandbox")  # Recommended for some environments
    # options.add_argument("--disable-dev-shm-usage")
    options.add_argument("user-data-dir=C:/Users/user/AppData/Local/Google/Chrome/User Data")
    browser = uc.Chrome(options=options)
    try:

        # Navigate to the PageSpeed Insights URL
        browser.get("https://pagespeed.web.dev/")
        print("On the site now...")
        # Wait for the search bar to load
        wait = WebDriverWait(browser, 10)
        search_bar = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[placeholder='Enter a web page URL']")))
        # Enter the target URL into the search bar
        search_bar.send_keys(company_url)
        print("URL entered in the search bar.")
        # Trigger the search (send Enter key)
        search_bar.send_keys(Keys.ENTER)
        print("Search initiated...")
        # Wait for results to load
        time.sleep(30)
        time.sleep(40)  # Adjust as needed for results to load
        # Mobile Speed Index
        print("40 seconds Completed")
        time.sleep(30)
        print("70 secs completed!!")
        time.sleep(40)
        print("110 secs completed, bro!")
        speed_index_element = wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="speed-index"]/div/div[2]')))
        speed_index_value = speed_index_element.text
        print(f"Mobile Speed Index: {speed_index_value}")
        # Mobile Friendliness Score:
        MobileFriendlinesScore = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[2]/span/div/div[2]/div[2]/div/div/article/div/div[2]/div/div/div/div[2]/a[2]/div[2]'))
        )
        MobileFriendlinesScore_value = MobileFriendlinesScore.text
        print(f"MFS: {MobileFriendlinesScore_value}")

        P_S = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[2]/span/div/div[2]/div[2]/div/div/article/div/div[2]/div/div/div/div[2]/a[1]/div[2]'))
        )
        Page_Strength = P_S.text
        print(f"Page Strenght: {Page_Strength}")
        
        FCP = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[2]/span/div/div[2]/div[2]/div/div/article/div/div[3]/div[2]/div[1]/div/div[2]/div[2]/div[1]/div/div[2]'))
        )
        FCPvalue = FCP.text
        print(f"FCP Speed Index: {FCPvalue}")
        LCP = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[2]/span/div/div[2]/div[2]/div/div/article/div/div[3]/div[2]/div[1]/div/div[2]/div[2]/div[2]/div/div[2]'))
        )
        LCPvalue = LCP.text
        print(f"LCP Index: {LCPvalue}")
        TBT = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[2]/span/div/div[2]/div[2]/div/div/article/div/div[3]/div[2]/div[1]/div/div[2]/div[2]/div[3]/div/div[2]'))
        )
        TBTvalue = TBT.text
        print(f"TBT Speed Index: {TBTvalue}")
        CLS = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[2]/span/div/div[2]/div[2]/div/div/article/div/div[3]/div[2]/div[1]/div/div[2]/div[2]/div[4]/div/div[2]'))
        )
        CLSvalue = CLS.text
        print(f"CLS Speed Index: {CLSvalue}")
        
        # Switch to Desktop tab
        print("Waiting to click on the Desktop to collect page speed")
        desktop_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="desktop_tab"]'))
        )
        # Scroll into view and click
        browser.execute_script("arguments[0].scrollIntoView(true);", desktop_tab)
        browser.execute_script("arguments[0].click();", desktop_tab)
        print("Clicked on the Desktop tab...")
        time.sleep(10)

        desktop_speed_index = wait.until(
            EC.presence_of_element_located((By.XPATH, '/html/body/c-wiz/div[2]/div/div[2]/div[3]/div/div/div[3]/span/div/div[2]/div[2]/div/div/article/div/div[3]/div[2]/div[1]/div/div[2]/div[2]/div[5]/div/div[2]'))
        )
        speed_index_value_desktop = desktop_speed_index.text
        print(f"Desktop Speed Index: {speed_index_value_desktop}")
        # Return both values
        return speed_index_value, speed_index_value_desktop, MobileFriendlinesScore_value, FCPvalue, LCPvalue, TBTvalue, CLSvalue, Page_Strength
    except TimeoutException:
        print("Page load timed out or element not found.")
        return None, None, None, None, None, None, None, None
    except NoSuchElementException:
        print("Required element not found on the page.")
        return None, None, None, None, None, None, None, None
    finally:
        # Close the browser
        print("Load time completed!... quitting browser")
        browser.quit()

# Function to scrape the Speed Index from PageSpeed Insights
def domain_rating(company_url):
    # Setup undetected ChromeDriver
    options = uc.ChromeOptions()
    # options.add_argument("--headless=new")
    # options.add_argument("--no-sandbox")  # Recommended for some environments
    # options.add_argument("--disable-dev-shm-usage")
    options.add_argument("user-data-dir=C:/Users/user/AppData/Local/Google/Chrome/User Data")
    browser = uc.Chrome(options=options)
    try:
        # Navigate to the PageSpeed Insights URL
        browser.get("https://moz.com/domain-analysis")
        print("On the site now...")
        # Wait for the search bar to load
        wait = WebDriverWait(browser, 10)
        search_bar = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[placeholder='Enter a domain']")))
        
        # Enter the target URL into the search bar
        search_bar.send_keys(company_url)
        print("URL entered in the search bar.")
        # Trigger the search (send Enter key)
        search_bar.send_keys(Keys.ENTER)
        print("Search initiated...")
        # Wait for results to load
        time.sleep(15)
        
        DomainAuthority = wait.until(EC.presence_of_element_located((By.XPATH, '/html/body/div[2]/main/div[2]/div/div/div/div/div/div[2]/h1')))
        DomainAuthority_value = DomainAuthority.text
        print(f"Domain Authority Value: {DomainAuthority_value}")
        return DomainAuthority_value
    except TimeoutException:
        print("Page load timed out or element not found.")
        return None, None
    except NoSuchElementException:
        print("Required element not found on the page.")
        return None, None
    finally:
        # Close the browser
        print("Load time completed!... quitting browser")
        browser.quit()


def convert_to_grade(score):
    score = float(score)  # Convert the string to a float (or int if preferred)
    
    if 90 <= score <= 100:
        return 'A+'
    elif 80 <= score < 90:
        return 'A'
    elif 70 <= score < 80:
        return 'B+'
    elif 60 <= score < 70:
        return 'B'
    elif 50 <= score < 60:
        return 'C+'
    elif 40 <= score < 50:
        return 'D'
    elif 30 <= score < 40:
        return 'E'
    else:
        return 'F'
def count_numbers_less_than_10(arr):
    count = 0
    for num in arr:
        if num < 10:
            count += 1
    return count

# Function to scroll to the bottom of the page
def scroll_to_bottom(browser):
    """Scrolls to the bottom of the page."""
    print("Scrolling to the bottom of the page...")
    browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(2)  # Allow time for dynamic content to load

# Function to scroll an element into view
def scroll_to_element(element, browser):
    """Scrolls the element into view."""
    print("Scrolling to the 'Show More' button...")
    browser.execute_script("arguments[0].scrollIntoView(true);", element)
    time.sleep(1)  # Pause after scrolling


EMAIL = "dm@eplanetsoft.com"
PASSWORD = "Eplanet@DM*&1234!"
EXPECTED_URL = "https://app.ahrefs.com/dashboard"

def scrape_data(company_name, company_url, company_xpath):
    options = uc.ChromeOptions()
    # options.add_argument("--headless=new")
    # options.add_argument("--no-sandbox")  # Recommended for some environments
    # options.add_argument("--disable-dev-shm-usage")
    # options.add_argument("--log-level=3")  # Reduces unnecessary logs
    # options.add_argument("--enable-logging")  # Enables logging
    options.add_argument("user-data-dir=C:/Users/user/AppData/Local/Google/Chrome/User Data")
    browser = uc.Chrome(options=options)

    try:
        # Step 1: Navigate to the website
        browser.get("https://ahrefs.com/")
        print("Navigated to Ahrefs homepage.")
        time.sleep(5)
        print("Okay!")

        # Step 2: Click on Sign In or Login button
        login_button = WebDriverWait(browser, 10).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Sign in"))
        )
        login_button.click()
        print("Clicked on Sign In.")
        time.sleep(2)
        if browser.current_url != EXPECTED_URL:
            google_button = WebDriverWait(browser, 10).until(
                EC.element_to_be_clickable((By.XPATH, "/html/body/div[1]/div/div[2]/div/div/div/div/div[1]/div/div/div/div[1]/button"))
            )
            google_button.click()
            time.sleep(3)

                # Switch to popup window
            browser.switch_to.window(browser.window_handles[-1])
            time.sleep(2)
                
            try:
                    # Check if EMAIL is available
                account = browser.find_element(By.XPATH, f"//div[text()='{EMAIL}']")
                time.sleep(1)
                account.click()
            except:
                    # Click 'Use another account'
                add_account = WebDriverWait(browser, 5).until(
                    EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Use another account')]"))
                )
                time.sleep(2)
                add_account.click()
                time.sleep(1)
                email_input = WebDriverWait(browser, 10).until(
                    EC.presence_of_element_located((By.ID, "identifierId"))
                )
                email_input.send_keys(EMAIL)
                browser.find_element(By.ID, "identifierNext").click()
                time.sleep(2)
                    
                    # Enter password and click next
                password_input = WebDriverWait(browser, 10).until(
                    EC.presence_of_element_located((By.NAME, "Passwd"))
                )
                checkbox = WebDriverWait(browser, 10).until(
                    EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/div[1]/div[2]/c-wiz/div/div[2]/div/div/div[1]/form/span/section[2]/div/div/div[1]/div[3]/div/div[1]/div/div/div[1]/div/div/input"))
                )
                time.sleep(1)
                checkbox.click()
                time.sleep(2)
                password_input.send_keys(PASSWORD)
                time.sleep(2)
                browser.find_element(By.ID, "passwordNext").click()
                time.sleep(3)
                
            # Wait for 2FA approval
            time.sleep(60)
            # Switch back to main window
            browser.switch_to.window(browser.window_handles[0])
            time.sleep(3)
            
            if browser.current_url == EXPECTED_URL:
                print("Login successful!")
            else:
                print("Login failed! Check your credentials or 2FA.")
                browser.quit()
        
            
        company_data = {}

        company_data[company_name] = {
            "XPATH": company_xpath,
            "URL": company_url 
        }

        print("Company data stored successfully!")
        print(company_data)

        SiteExplorer = WebDriverWait(browser, 10).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Site Explorer"))
        )
        SiteExplorer.click()
        print("Clicked on Site Explorer.")
        time.sleep(10)

        
        wait = WebDriverWait(browser, 10)
        search_bar = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[placeholder='Domain or URL']")))
            
        # Enter the target URL into the search bar
        search_bar.send_keys(company_url)
        print("URL entered in the search bar.")

        time.sleep(5)

        Search = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div/div[2]/form/div/button'))
        )
        Search.click()
        print("Search bar pressed")

        Backlinks = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, 'Backlinks'))
        )
        Backlinks.click()
        print("Clicked on the Backlinks...")

        time.sleep(3)

        Dofllow_Link = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[1]/div[2]/div/div/div/div[1]/div[2]/button/div/div'))
        )
        Dofllow_Link.click()
        print("Clicked on the Dofollow Links....")

        time.sleep(2)

        Dofllow_Link_show = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[1]/div[2]/div/div[2]/div/button[1]/div'))
        )
        Dofllow_Link_show.click()
        print("Clicked on the Show results....")

        time.sleep(4)

        DofollowLinks_value  = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[1]/div[3]/div/div/div[1]/div/div/div/div[1]/div/div'))
        ).text
        print(f"Dofollow Links: {DofollowLinks_value}")

        time.sleep(3)

        Nofllow_Link = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[1]/div[2]/div/div/div/div[1]/div[3]/div/div/div/button'))
        )
        Nofllow_Link.click()
        print("Clicked on the Nofollow Links....")
        time.sleep(2)

        nofollow_option = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[11]/div/div/div/div/div/div[1]/button'))
        )
        nofollow_option.click()
        print("Clicked on the Nofollow option")
        
        Nofllow_Link_show = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[1]/div[2]/div/div[2]/div/button[1]/div'))
        )
        Nofllow_Link_show.click()
        print("Clicked on the Show results....")

        time.sleep(4)

        NofollowLinks_value  = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[1]/div[3]/div/div/div[1]/div/div/div/div[1]/div/div'))
        ).text
        print(f"Nofollow Links: {NofollowLinks_value}")

        Site_audit = WebDriverWait(browser, 10).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Site Audit"))
        )
        Site_audit.click()
        print("Clicked on Site Audit.")
        time.sleep(5)


        # This is to select the company.... change..
        print(f"{company_name} selected!")
        Company_to_click = WebDriverWait(browser, 20).until(
            EC.element_to_be_clickable((By.XPATH, company_xpath))
        )
        
        # Click the element
        Company_to_click.click()
        print(f"Company url: {company_url}")
        
        print("Gettting thee real info now")


        # Step 5: Navigate to Internal Link Opportunities
        Internal_Link_Opportunity = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, 'Internal link opportunities'))
        )
        Internal_Link_Opportunity.click()
        print("Clicked on the Internal Link Opportunity...")

        # Initialize variables
        wait = WebDriverWait(browser, 10)
        pr_values = []  # Final array to store collected PR values
        show_more_xpath = '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[2]/div/div[2]/div[2]/div[4]/div[1]'  # XPath for 'Show More'

        try:
            while True:
                # Step 6: Collect PR values
                temp_pr_values = []
                try:
                    print("Collecting PR values...")
                    row_elements = wait.until(EC.presence_of_all_elements_located(
                        (By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[2]/div/div[2]/div[2]/div[3]/div/div[2]/div/div[1]/div/div/table/tbody[2]/tr'))
                    )

                    for row in row_elements:
                        try:
                            pr_element = row.find_element(By.XPATH, './td[2]')
                            pr_text = pr_element.text.strip()
                            if pr_text.isdigit():
                                temp_pr_values.append(int(pr_text))
                                print(f"Collected PR: {pr_text}")
                        except Exception as e:
                            print(f"Could not extract PR: {e}")

                    pr_values = temp_pr_values  # Replace the final array with the current iteration values

                except TimeoutException:
                    print("No more rows found or timeout occurred.")
                    break

                # Step 7: Scroll and click 'Show More' button
                try:
                    scroll_to_bottom(browser)  # Scroll to bottom before searching for the button
                    show_more_button = browser.find_element(By.XPATH, show_more_xpath)

                    if show_more_button.is_displayed() and show_more_button.is_enabled():
                        scroll_to_element(show_more_button, browser)  # Bring button into view
                        try:
                            show_more_button.click()
                            print("Clicked 'Show More' button.")
                            time.sleep(2)  # Wait for rows to load
                        except Exception as e:
                            print(f"Direct click failed, trying ActionChains: {e}")
                            ActionChains(browser).move_to_element(show_more_button).click().perform()
                            time.sleep(2)
                    else:
                        print("'Show More' button not interactable. Exiting loop.")
                        break
                except NoSuchElementException:
                    print("No 'Show More' button found. Ending collection.")
                    break

        except Exception as e:
            print(f"An error occurred: {e}")

        finally:
            print("Final PR Values:", pr_values)

        
        # Step 1: Navigate to Overview for Health Score
        overview_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Overview"))
        )
        overview_tab.click()

        print("Clicked on the Overview...")

        # Wait for Health Score element
        try: 
            health_score = WebDriverWait(browser, 30).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div[1]/div[2]/div/div/div/div[2]/div/div[1]/div[2]/div/div[1]'))
            ).text
            print(f"Health Score: {health_score}")
        except:
            print("Error in finding health score need to change the Xpath again")

        Health_word = WebDriverWait(browser, 30).until(
        EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div[1]/div[2]/div/div/div/div[2]/div/div[1]/div[2]/div/div[2]/div/div'))
        ).text
        print(f"Health Score Word: {Health_word}")

        # Step 2: Navigate to Content for Meta Descriptions
        content_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Content"))
        )
        content_tab.click()

        title_tag_section = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div/div/div[2]/div[1]/div'))
        )

        title_tag = title_tag_section.find_elements(By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div/div/div[2]/div[1]/div/div/div[2]/div/div[2]/div')

        for value in title_tag:
            print(f"Title Tag setup: {value.text}")

        title_tag_link = [value.text for value in title_tag]

        
        # Scrape Meta Descriptions
        # Wait for the "Meta description tag setup" section to load
        meta_description_section = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div/div/div[3]/div[1]'))
        )
        
        # Locate and extract the values under the section
        meta_issues_values = meta_description_section.find_elements(By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div/div/div[3]/div[1]/div/div/div[2]/div/div[2]/div')  # Adjust the class name if needed
        
        # Print the extracted values
        for value in meta_issues_values:
            print(f"Meta Description Issue: {value.text}")

        meta_description_issues = [value.text for value in meta_issues_values]

        # Scrape Meta Descriptions
        # Wait for the "Meta description tag setup" section to load
        H1_setup = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div/div/div[4]/div[1]'))
        )
        
        # Locate and extract the values under the section
        H1_values = H1_setup.find_elements(By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div/div/div[4]/div[1]/div/div/div[2]/div/div[2]/div')  # Adjust the class name if needed
        
        # Print the extracted values
        for value in H1_values:
            print(f"H1 Tag description: {value.text}")

        H1_tag_issues = [value.text for value in H1_values]


        # Step 3: Navigate to Links for Broken Links and Internal Redirect Links
        links_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Links"))
        )
        links_tab.click()

        print("Navigated to Links sections now....")

        # Scrape Broken Links
        broken_internal_link = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[1]/div[2]/div/div/div[2]'))
        ).text
        print(f"Broken internal Links: {broken_internal_link}")

        broken_external_link = WebDriverWait(browser, 20).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[1]/div[4]/div/div/div[2]'))
        ).text
        print(f"Broken External Links: {broken_external_link}")


        Internal_Links = WebDriverWait(browser, 20).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[1]/div[1]/div/div/div[2]/a/div'))
        ).text
        print(f"Internal Links: {Internal_Links}")

        

        # Step 4: Navigate to Bulk Export for Redirect Chains and Dofollow/Nofollow Links
        bulk_export_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Bulk export"))
        )
        bulk_export_tab.click()

        time.sleep(2)
        # Orphan pages basically...
        Internal_Link_Audit  = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[8]/td[4]'))
        ).text
        print(f"Internal Link Audit: {Internal_Link_Audit}")

        # Robots_txt = WebDriverWait(browser, 20).until(
        #     EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[20]/td[4]'))
        # ).text
        # print(f"Robots.txt link blocked config: {}")
        # Scrape Redirect Chains
        redirect_chains = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[7]/td[4]'))
        ).text
        print(f"Redirect Chains: {redirect_chains}")

        Internal301link = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[22]/td[4]'))
        ).text
        print(f"Internal 301 link blocks: {Internal301link}")

        Duplicate_content = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[6]/td[4]'))
        ).text
        print(f"Duplicate Content: {Duplicate_content}")

        MixedContent = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[17]/td[4]'))
        ).text
        print(f"Mixed Content: {MixedContent}")


        Image_references_without_all_text = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[16]/td[4]'))
        ).text

        print(f"Image Optimization: {Image_references_without_all_text}")

        dofollow_links = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[29]/td[4]'))
        ).text
        nofollow_links = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[30]/td[4]'))
        ).text
        print(f"Dofollow Links: {dofollow_links}")
        print(f"Nofollow Links: {nofollow_links}")

        
        Custom401Page = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/table/tbody/tr[23]/td[4]'))
        ).text
        print(f"Custom 401 Page: {Custom401Page}")

        #navigating to performace tab
        performance_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Performance"))
        )
        performance_tab.click()

        Device = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[2]/div[1]/div/div/div[2]'))
        ).text
        print(f"Device: {Device}")

        Mobile_Friendliness_Score = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[2]/div[4]/div/div/div[2]'))
        ).text
        print(f"Mobile Friendliness Score: {Mobile_Friendliness_Score}")
        
        Mobile_friendliness_score_outof_total = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[1]/div[1]/div[2]/div/div/div[2]/div/div[2]/div'))
        ).text
        
        # Print the scraped text for debugging
        print(f"Out of total friendliness score: {Mobile_friendliness_score_outof_total}")
        
        # Example scraped text from the image you shared
        # Mobile_friendliness_score_outof_total might look like this:
        # "Needs improvement 24\nGood 16\nPoor 2"
        
        # Extract numbers from the text using regex
        numbers = map(int, re.findall(r'\d+', Mobile_friendliness_score_outof_total))
        
        # Calculate the total number of pages
        total_pages = sum(numbers)
        
        # Print the total pages
        print("Total Pages:", total_pages)

        TTFBD = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[3]/div[1]/div[1]'))
        )
        
        # Locate and extract the values under the section
        TTF = TTFBD.find_elements(By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[3]/div[1]/div[1]/div/div/div[2]/div/div[2]/div')  # Adjust the class name if needed
        
        # Print the extracted values
        for value in TTF:
            print(f"Time to First Byte Distribution: {value.text}")

        TTFBDIS = [value.text for value in TTF]

        LT = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[3]/div[1]/div[2]'))
        )
        
        # Locate and" extract the values under the section
        load_time = LT.find_elements(By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[3]/div[1]/div[2]/div/div/div[2]/div/div[2]/div')  # Adjust the class name if needed
        
        # Print the extracted values
        for value in load_time:
            print(f"Load Time Distribution: {value.text}")

        LTD = [value.text for value in load_time]

        File = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[3]/div[2]/div[1]'))
        )
        
        # Locate and extract the values under the section
        File_size = File.find_elements(By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[3]/div[3]/div[2]/div[1]/div/div/div[2]/div/div[2]/div')  # Adjust the class name if needed
        
        # Print the extracted values
        for value in File_size:
            print(f"File Size Distribution: {value.text}")

        FSD = [value.text for value in File_size]

        RedirectTab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Redirects"))
        )
        RedirectTab.click()

        try:
            # Try to locate the sitemap element
            Sitemap_element = WebDriverWait(browser, 30).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[2]/div/div[5]/div/div/div/div[2]/div/div[2]/div/div[1]/div/div/table/tbody/tr[9]/td[3]'))
            )
            # If found, get the text content
            Sitemap = Sitemap_element.text
            print(f"Sitemap.xml links: {Sitemap}")
        except Exception as e:
            # If not found, set Sitemap to 'N/A'
            Sitemap = 'N/A'
            print("Sitemap.xml not present. Setting Sitemap to 'N/A'.")

        Indexability_Tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Indexability"))
        )
        Indexability_Tab.click()

        Indexed_Pages = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[1]/div[2]/div/div/div[2]/a/div'))
        ).text
        print(f"Indexed Pages: {Indexed_Pages}")

        TotalPages = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[1]/div[1]/div/div/div[2]/a/div'))
        ).text
        print(f"Total Pages available for indexed pages {TotalPages}")
        
        Total_Pages = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[1]/div[1]/div/div/div[2]/a/div'))
        ).text
        print(f"Total pages: {Total_Pages}")


        InternalPagesTab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Internal pages"))
        )
        InternalPagesTab.click()

        # collects the number of HTML Pages basically...
        # URL_Structure = WebDriverWait(browser, 30).until(
        #     EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[3]/div/div/div[3]/div/div/div[2]/div/div[1]/div[2]/div/div/div[2]/a/div'))
        # ).text
        # print(f"URL Structure: {URL_Structure}")

        # Wait for the HTTPS Status codes distribution section to load
        sslcertificate_httpstatuscode = WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[2]/div/div[1]/div[2]'))
        )
        
        # Locate and extract the values under the section
        sslcertificate = sslcertificate_httpstatuscode.find_elements(By.XPATH, '//*[@id="body"]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[2]/div/div[1]/div[2]/div/div/div[2]/div/div[2]/div')  # Adjust the class name if needed
        
        # Print the extracted values
        for value in sslcertificate:
            print(f"SSL Certificate: {value.text}")

        SSL_Certificate = [value.text for value in sslcertificate]
        

        Images_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Images"))
        )
        Images_tab.click()

        print("Navigated to Images sections now....")

        issue_tab = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[1]/button[2]/div/div/div[1]'))
        )
        issue_tab.click()

        print("Navigated to issue_tab  now....")
        # Locate the table rows (adjust the XPath or CSS selectors as needed)

        ImageIssues = WebDriverWait(browser, 20).until(
            EC.presence_of_element_located((By.XPATH, '/html/body/div[1]/div/div[2]/div[1]/div[2]/div/div/div[3]/div/div/div[2]/div/div[2]/div/div/table/tbody'))
        ).text
        print(f"Image Issue: {ImageIssues}")


    except Exception as e:
        print("An error occurred:", str(e))

    finally:
        browser.quit()


    robots_status, sitemap_status = check_robots_txt_and_sitemap(company_url)
    # Print results
    print(robots_status)
    print(sitemap_status)


    # Example usage
    mobile_speed, desktop_speed, MobileFS, FCP, LCP, TBT, CLS, Page_StrengthS = get_speed_index(company_url)

    if mobile_speed and desktop_speed and MobileFS and FCP and LCP and TBT and CLS and Page_StrengthS:
        print("cool!")
    else:
        print("\nFailed to retrieve speed index.")

    domainAuthority = domain_rating(company_url)

    if domainAuthority:
        print(f"Domain Authority Value: {domainAuthority}")
    else:
        print("\nFailed to retrieve speed index.")


    grade = convert_to_grade(health_score)
    print(f"Score: {health_score}, Grade: {grade}")

    import pandas as pd
    # Create a DataFrame with the extracted data
    data = {
        'Domain Strength/Authority': domainAuthority,
        'Meta Description Issues': meta_description_issues,
        'H1 Tag Issues': H1_tag_issues,
        'Time to first byte distribution': TTFBDIS,
        'Load Time distribution': LTD,
        'File Size Distribution': FSD,
        'Health Score': health_score,
        'Broken Internal Links': broken_internal_link,
        'Broken External Links': broken_external_link,
        'Internal Links': Internal_Links,
        'Internal Link Audit': Internal_Link_Audit,
        'Redirect Chains': redirect_chains,
        'Robots.txt Blocks': robots_status,
        'Internal 301 Link Blocks': Internal301link,
        'Duplicate Content': Duplicate_content,
        'Mixed Content': MixedContent,
        'Dofollow Links': DofollowLinks_value,
        'Nofollow Links': NofollowLinks_value,
        'Custom 401 Page': Custom401Page,
        'Device Rendering': Device,
        'Mobile Friendliness Score': MobileFS, 
        'Sitemap.xml link': sitemap_status,
        'Indexed Pages': Indexed_Pages,
        'SSL Certificate': SSL_Certificate,
        'Image Optimization': Image_references_without_all_text,
        'Page Strength': Page_StrengthS,
        'Health_Score_word': Health_word,
        'Title_tags': title_tag_link,
        'TotalPages': TotalPages,
        'SpeedIndexMobile': mobile_speed,
        'SpeedIndexDesktop':desktop_speed,
        'UsabilityGrade': grade,
        'Image Issues': ImageIssues,
        'First Contentful Paint':FCP, 
        'Largest Contentful Paint':LCP, 
        'Total Blocking Time':TBT, 
        'Cumulatice Layout Shift':CLS
    }

    # Convert all non-list values to lists
    for key in data:
        if not isinstance(data[key], list):
            data[key] = [data[key]]  # Convert single values or strings to single-item lists

    # Determine the maximum length of the arrays in the dictionary
    max_length = max(len(v) for v in data.values())

    # Adjust all lists to he same length by padding with None
    for key in data:
        if len(data[key]) < max_length:
            data[key].extend([None] * (max_length - len(data[key])))

    # Convert the data into a pandas DataFrame
    df = pd.DataFrame(data)
    print("Saving the information!")
    # Save the DataFrame to a CSV file
    
    # Sanitize the company_url to remove invalid characters for file names
    sanitized_url = re.sub(r'[<>:"/\\|?*]', '_', company_url)
    # Save the DataFrame to a CSV file with the sanitized name
    csv_file = f'Scraped_Data_{sanitized_url}.csv'
    df.to_csv(csv_file, index=False)
    print("File saved!")
    # Optional: Display the DataFrame to verify before saving
    print(df.head())
    return csv_file


def get_chatgpt_response(prompt):
    
    openai_api_key = os.getenv("API_KEY")

    client = OpenAI(api_key=openai_api_key)
    try:
        messages = [
            {"role": "system", "content": "You are a helpful assistant for creating SEO reports."},
            {"role": "user", "content": prompt}
        ]
        # Updated API call
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=200,
            temperature=0
        )
        return response.choices[0].message.content  # Accessing response content
    except Exception as e:
        return f"Error fetching response from ChatGPT: {str(e)}"

# Function to add headings
def add_heading(doc, text, level, color=(0, 0, 0)):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.font.name = 'Calibri'  # Set the font to Calibri
    heading_run.font.italic = False
    heading_run.font.color.rgb = RGBColor(*color)

# Function to add a bullet list
def add_bullet(doc, text, level=1):
    # Use the default bullet list style without specifying the level
    para = doc.add_paragraph(text, style='List Bullet')
    
    # Indent bullets based on the level
    para.paragraph_format.left_indent = Inches(0.3 * level)
    run = para.runs[0]
    run.font.name = 'Calibri'

# Function to add a bullet point with specific bold text
def add_bullet_bold(doc, text, level=1):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.3 * level)
    
    # Add bold part for the specific text
    run = para.add_run(text)
    run.bold = True  # Make the text bold


def generate_seo_report(csv_file, company_name):        

    # Load CSV data
    try:
        df = pd.read_csv(csv_file)
    except FileNotFoundError:
        print("Error: CSV file not found. Please check the file name and path.")
        exit()

    # Create a new Word document
    doc = Document()



    # Add Title
    add_heading(doc, f'{company_name} - SEO Audit Report', 1)

    # Add Introduction Section
    add_heading(doc, 'Introduction', level=2)
    intro_paragraph = doc.add_paragraph()
    intro_content = (
        f"This report provides a detailed analysis of the SEO health of {company_name}. "
        "The findings are based on a comprehensive audit and inspection. "
        "This report outlines the key issues affecting the website's SEO performance, their potential impact, and the importance of addressing these issues."
    )

    intro_run = intro_paragraph.add_run(intro_content)
    intro_run.font.name = 'Calibri'  # Set font to Calibri
    intro_run.font.size = Pt(11)  # Optionally set font size


    # Overview of SEO Health
    add_heading(doc, '1. Overview of SEO Health', level=3)
    # health_score = df["Health Score"].iloc[0] if "Health Score" in df.columns else "N/A"
    overall_health_score = doc.add_paragraph(style='List Bullet 3')
    overall_health_score.add_run('Overall Health Score:').bold = True
    overall_health_score.add_run(f' {df["Health_Score_word"].iloc[0]}')
    overall_health_score.paragraph_format.left_indent = Inches(0.7)

    # Add Key Metrics
    key_metrics = doc.add_paragraph(style='List Bullet 3')
    key_metrics.add_run('Key Metrics:').bold = True
    key_metrics.paragraph_format.left_indent = Inches(0.7)
    if "Domain Strength/Authority" in df.columns:
        add_bullet(doc, f"Domain Authority: {df['Domain Strength/Authority'].iloc[0]}", level=4)
    if "Load Time distribution" in df.columns:
        add_bullet(doc, f"Page Speed: Average of  {df['SpeedIndexDesktop'].iloc[0]} (Desktop), {df['SpeedIndexMobile'].iloc[0]} (Mobile)", level=4)
    if "Mobile Friendliness Score" in df.columns:
        add_bullet(doc, f"Mobile Friendliness: Score of {df['Mobile Friendliness Score'].iloc[0]}/100", level=4)
    if "Indexed Pages" in df.columns:
        add_bullet(doc, f"Indexed Pages (as per Ahrefs/Screaming Frog): {df['Indexed Pages'].iloc[0]} out of {df["TotalPages"].iloc[0]}", level=4)

    # Add Technical SEO Issues Section
    add_heading(doc, '2. Detailed Issues and Their Impact on SEO', level=3)
    add_heading(doc, '2.1. Technical SEO Issues', level=3)

    # Example: Crawlability and Indexability
    add_heading(doc, '2.1.1. Crawlability and Indexability', level=4)

    # Add Robots.txt Configuration Bullet Point with bold text
    add_bullet_bold(doc, 'Robots.txt Configuration:', level=2)

    Robot_Prompt = (
        f"Analyze whether the presence or absence of a Robots.txt file is an issue for SEO. "
        f"Data: Robots.txt Presence - {df['Robots.txt Blocks'].iloc[0]}. "
        f"If Robots.txt is missing, explain why this could be problematic for controlling crawler access. "
        f"If it exists, determine whether it blocks important pages. "
        f"Summarize the issue in a single line, such as: 'No Robots.txt file detected, meaning search engines have unrestricted access, which may lead to indexing of unnecessary pages.' "
        f"Rephrase the issue accordingly based on the given data."
    )

    Robot_Prompt1 = get_chatgpt_response(Robot_Prompt)

    # Add a bullet manually and get the paragraph object
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {Robot_Prompt1}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph




    # Add the rest of the text

    Impact_Prompt = (
        f"Impact: [Explain the impact, e.g., 'Blocking important pages from being crawled by search engines can severely limit the visibility of these pages in search results, reducing organic traffic.in a single line only]{df['Robots.txt Blocks'].iloc[0]}, using different wordings, also mention the data i have shared with you"
    )
    Impact_Prompt = get_chatgpt_response(Impact_Prompt)

    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{Impact_Prompt}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    add_bullet_bold(doc, 'Sitemap.xml: ', level=2)
    Sitemap_prompt_issue = (
        f"Summarize the issue in a single line, avoiding the use of 'Issue:'. Focus on how the Sitemap.xml is missing recently added pages, rephrasing creatively. Ensure to include the Sitemap link data: {df['Sitemap.xml link'].iloc[0]}."
    )
    Sitemap_prompt = get_chatgpt_response(Sitemap_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {Sitemap_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    Impact_Prompt = (
        f"Impact: [Explain the impact, e.g., 'If important pages are not included in the sitemap, search engines may not discover or index these pages, leading to a loss in potential rankings.'] in a single line{df['Sitemap.xml link'].iloc[0]}using different wordings, also mention the data i have shared with you"
    )
    Impact_Prompt = get_chatgpt_response(Impact_Prompt)

    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{Impact_Prompt}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph




    add_heading(doc, '2.1.2. HTTPS/SSL Implementation', level=4)


    add_bullet_bold(doc, 'SSL Certificate: ', level=2)

    SSL_Issue_issue = (
        f"Based on the provided data, summarize in a single line why certain pages might still be served over HTTP despite the SSL certificate being installed correctly. Ensure the explanation creatively incorporates the SSL certificate details: {df['SSL Certificate'].iloc[0]}."
    )

    SSL_prompt = get_chatgpt_response(SSL_Issue_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {SSL_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    Impact_Prompt = (
        f'Summarize in the single line, how Serving pages over HTTP instead of HTTPS can lead to security warnings in browsers, which may deter users and negatively affect the site trustworthiness in search engines, {df['SSL Certificate'].iloc[0]} using different wordings'
    )
    Impact_Prompt = get_chatgpt_response(Impact_Prompt)

    # Assuming 'doc' is your Document object
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{Impact_Prompt}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    add_bullet_bold(doc, 'Mixed Content: ', level=2)

    Mixed_content_issue = (
        f"Summarize in a single line why certain pages have mixed content issues, highlighting that images are being loaded over HTTP. Incorporate the provided mixed content data dynamically: {df['Mixed Content'].iloc[0]}."
    )

    MixedContent_prompt = get_chatgpt_response(Mixed_content_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {MixedContent_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    Impact_Prompt = (
        f'Summarize  the impact in a single line e.g., "Mixed content can lead to security warnings and reduced user trust, as well as potential SEO penalties from search engines prioritizing secure sites." in a single line, {df['SSL Certificate'].iloc[0]}, using different wordings'
    )
    Impact_Prompt = get_chatgpt_response(Impact_Prompt)

    # Assuming 'doc' is your Document object
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{Impact_Prompt}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    add_heading(doc, '2.1.3. 404 Error Pages', level=4)


    add_bullet_bold(doc, 'Broken Links: ', level=2)

    BrokenLink_prompt_issue = (
        f"Summarize in a single line the issue of broken links on the site, mentioning both internal and external links. Include the following data: internal broken links: {df['Broken Internal Links'].iloc[0]}, external broken links: {df['Broken External Links'].iloc[0]}."
    )

    BrokenLink_prompt = get_chatgpt_response(BrokenLink_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {BrokenLink_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    Impact_Prompt = (
        f'Impact: [Explain the impact, e.g., "Broken links can create a poor user experience, leading to increased bounce rates. They also waste crawl budget, which could be better spent on indexing valuable pages."] in a single line, using different wordings, also mention the data i have shared with you'
    )
    Impact_Prompt = get_chatgpt_response(Impact_Prompt)

    # Assuming 'doc' is your Document object
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{Impact_Prompt}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    add_bullet_bold(doc, 'Custom 404 Page: ', level=2)

    Custom404_prompt_issue = (
        f"Summarize in one line the issue with the site's 404 page, highlighting missing navigation or customization. Use the provided data: {df['Custom 401 Page'].iloc[0]}. If the value is 0, state there is no issue."
    )

    Custom404 = get_chatgpt_response(Custom404_prompt_issue)

    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {Custom404}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    Impact_Prompt = (
        f'Impact: [Explain the impact, e.g., "A generic 404 page can result in users leaving the site, increasing bounce rates and reducing user engagement."]" in a single line, {df['Custom 401 Page'].iloc[0]}, using different wordings, also mention the data i have shared with you'
    )
    Impact_Prompt = get_chatgpt_response(Impact_Prompt)

    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{Impact_Prompt}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, '2.1.4. 301 Redirects', level=4)

    # Redirect Chains
    add_bullet_bold(doc, 'Redirect Chains: ', level=2)

    redirect_chains_prompt_issue = (
        f"Summarize in one line how redirect chains impact load times, using the provided data: {df['Redirect Chains'].iloc[0]}. If no chains exist (value is 0), state there is no issue."
    )

    redirect_chains_prompt = get_chatgpt_response(redirect_chains_prompt_issue)

    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {redirect_chains_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    redirect_chains_impact_prompt = (
        f"Summarize the impact, e.g., 'Redirect chains can slow down page load times, negatively affecting both user experience and search engine rankings.'] "
        f"In a single line, using different wording, also mention the data I have shared with you: '{df['Redirect Chains'].iloc[0]}'."
    )

    redirect_chains_impact = get_chatgpt_response(redirect_chains_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{redirect_chains_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Internal 301 Redirects
    add_bullet_bold(doc, 'Internal 301 Redirects: ', level=2)

    internal_redirects_prompt_issue = (
        f"Summarize in a single line how internal links are pointing to URLs with 301 redirects, potentially impacting site performance. Include the provided data: {df['Internal 301 Link Blocks'].iloc[0]}, and rephrase creatively."
    )

    internal_redirects_prompt = get_chatgpt_response(internal_redirects_prompt_issue)

    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {internal_redirects_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    internal_redirects_impact_prompt = (
        f" Summarize the impact, e.g., 'Using internal links that redirect can dilute link equity and slow down user navigation, potentially harming SEO performance.'] "
        f"In a single line, using different wording, also mention the data I have shared with you: '{df['Internal 301 Link Blocks'].iloc[0]}'."
    )

    internal_redirects_impact = get_chatgpt_response(internal_redirects_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{internal_redirects_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, '2.1.5. Duplicate Content', level=4)

    # Duplicate Pages
    add_bullet_bold(doc, 'Duplicate Pages: ', level=2)

    duplicate_pages_prompt_issue = (
        f"Summarize in a single line the issue of duplicate content on the site. Rephrase creatively and include the provided data: {df['Duplicate Content'].iloc[0]}."
    )

    duplicate_pages_prompt = get_chatgpt_response(duplicate_pages_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {duplicate_pages_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    duplicate_pages_impact_prompt = (
        f" Summarize in the single line the impact, e.g., 'Duplicate content can confuse search engines, leading to poor indexing and lower rankings due to content duplication penalties.'] "
        f"Based on the data '{df['Duplicate Content'].iloc[0]}', rephrase it in a single line using different wording and mention the data I have shared with you."
    )

    duplicate_pages_impact = get_chatgpt_response(duplicate_pages_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{duplicate_pages_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # URL Parameters
    add_bullet_bold(doc, 'URL Parameters: ', level=2)

    # Fetch relevant data from the dataset
    url_parameters_data = df[['Indexed Pages', 'Duplicate Content', 'Internal Links', 'Redirect Chains', 'Sitemap.xml link']].iloc[0]

    # Issue Analysis Prompt
    url_parameters_prompt_issue = (
        f"Analyze whether URL parameters in the given data are actually causing SEO issues. "
        f"Check if they create duplicate content, indexing problems, or unnecessary redirects. "
        f"Data: Indexed Pages - {url_parameters_data['Indexed Pages']}, "
        f"Duplicate Content - {url_parameters_data['Duplicate Content']}, "
        f"Internal Links - {url_parameters_data['Internal Links']}, "
        f"Redirect Chains - {url_parameters_data['Redirect Chains']}, "
        f"Sitemap.xml - {url_parameters_data['Sitemap.xml link']}. "
        f"If there's an issue, summarize it in one line. If there’s no major issue, explain why in one line."
    )

    url_parameters_prompt = get_chatgpt_response(url_parameters_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ').bold = True
    bullet_paragraph.add_run(f" {url_parameters_prompt}")
    bullet_paragraph.paragraph_format.left_indent = Inches(1)
    bullet_paragraph.style.font.name = 'Calibri'  



    url_parameters_impact_prompt = (
        f"Summarize in a single line the general impact of URL parameters on SEO. "
        f"Example: 'URL parameters that create duplicate pages can split link equity and dilute the effectiveness of content in search rankings.' "
        f"Provide a concise rephrasing of this impact based on general SEO knowledge."
    )

    url_parameters_impact = get_chatgpt_response(url_parameters_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{url_parameters_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph




    # Add heading
    add_heading(doc, '2.1.6. Site Speed Optimization', level=4)

    # Page Load Times
    add_bullet_bold(doc, 'Page Load Times: ', level=2)

    page_load_times_prompt_issue = (
        f"Summarize in one line how page load times exceed acceptable thresholds, possibly due to factors like large images. Include the load time data (in seconds): {df['Load Time distribution'].iloc[0]}. If it exists, then highlight that 'Fast' is good, and other values need improvement."
    )


    page_load_times_prompt = get_chatgpt_response(page_load_times_prompt_issue)

    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {page_load_times_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    page_load_times_impact_prompt = (
        f'Impact: [Explain the impact, e.g., "Slow page load times can lead to higher bounce rates, lower user engagement, and poor search engine rankings, especially on mobile devices."], in a single line, using different wordings'
    )
    page_load_times_impact = get_chatgpt_response(page_load_times_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{page_load_times_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Time to First Byte (TTFB)
    add_bullet_bold(doc, 'Time to First Byte (TTFB):', level=2)

    ttfb_prompt_issue = (
        f"Summarize in one line how the homepage's TTFB (Time to First Byte) exceeds recommended thresholds. Include the TTFB distribution data (in seconds): {df['Time to first byte distribution'].iloc[0]}. If it exists, then highlight that 'Fast' is good, while 'Medium,' 'Slow,' or 'Too Slow' require improvement."
    )


    ttfb_prompt = get_chatgpt_response(ttfb_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {ttfb_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    ttfb_impact_prompt = (
        f'Impact: [Explain the impact, e.g., "High TTFB can negatively impact page load speed, which is a known ranking factor in Google’s algorithm."], in a single line, using different wordings'
    )
    ttfb_impact = get_chatgpt_response(ttfb_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{ttfb_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Image Optimization
    add_bullet_bold(doc, 'Image Optimization:', level=2)

    image_issue_prompt = (
        f"Summarize in a single line the issue regarding images on the site, based on the following data: "
        f"'{df['Image Issues'].iloc[0]}' (e.g., large file sizes). "
        f"Rephrase creatively and indicate potential SEO impact. Also include the data!,if the data is empty or zero or null then say that no issues till now!"
    )

    image_optimization_prompt = get_chatgpt_response(image_issue_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {image_optimization_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    image_optimization_impact_prompt = (
        f'Impact: [Explain the impact, e.g., "Large, unoptimized images can significantly slow down page loading times, negatively affecting both user experience and SEO rankings."], in a single line, using different wordings'
    )
    image_optimization_impact = get_chatgpt_response(image_optimization_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{image_optimization_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, '2.1.7. XML Sitemap', level=4)

    # Sitemap Inclusion
    add_bullet_bold(doc, 'Sitemap Inclusion: ', level=2)

    sitemap_inclusion_prompt_issue = (
        f"Summarize in a single line how the XML sitemap fails to include the latest blog posts or newly added pages. Rephrase creatively and keep it concise. i also need you to, Include the data, {df["Sitemap.xml link"].iloc[0]}, if the data is empty or zero or null then say that no issues till now!"
    )

    sitemap_inclusion_prompt = get_chatgpt_response(sitemap_inclusion_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {sitemap_inclusion_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    sitemap_inclusion_impact_prompt = (
        '[Explain the general impact of the XML sitemap in terms of SEO, in a single line, using different wordings'
    )
    sitemap_inclusion_impact = get_chatgpt_response(sitemap_inclusion_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{sitemap_inclusion_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, '2.1.8. Internal Linking Structure', level=4)

    # Internal Link Audit
    add_bullet_bold(doc, 'Internal Link Audit: ', level=2)

    internal_link_audit_prompt_issue = (
        f"Summarize in a single line how some pages are orphaned with no internal links pointing to them. Include the internal link audit data: {df['Internal Link Audit'].iloc[0]}, and rephrase creatively. if the data is empty or zero or null then say that no issues till now!"
    )

    internal_link_audit_prompt = get_chatgpt_response(internal_link_audit_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {internal_link_audit_prompt}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    internal_link_audit_impact_prompt = (
        'Impact: [Explain the impact, e.g., "Orphan pages are difficult for search engines to find and index, leading to lower visibility in search results and missed opportunities for ranking."], in a single line, using different wordings'
    )
    internal_link_audit_impact = get_chatgpt_response(internal_link_audit_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{internal_link_audit_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Add heading
    add_heading(doc, '2.2. On-Page SEO Issues', level=3)


    # Add heading
    add_heading(doc, 'Title Tags', level=3)

    # Title Tags: Issue and Impact
    add_bullet_bold(doc, 'Title Tags: ', level=2)

    title_tags_prompt_issue = (
        f"Summarize in a single line how some pages have duplicate title tags, potentially affecting SEO. "
        f"Rephrase creatively and keep it concise. "
        f"Include the data: '{df['Title_tags'].iloc[0]}'. "
        f"Rephrase creatively. Please note that if the data contains only 'Only one' (even if followed by numbers) just mention that there is no issue with the title tags. and if and onnly if the data contains any value indicating more than one, mention that the title tags need improvement. if the data is empty or zero or null then say that no issues till now!"

    )


    title_tags_issue = get_chatgpt_response(title_tags_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {title_tags_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    title_tags_impact_prompt = (
        f"Summarize in a single line the impact of duplicate title tags. "
        f"Impact: [Explain the impact, e.g., 'Duplicate title tags can confuse search engines and reduce the effectiveness of keywords in improving rankings for individual pages.'] using different wordings. "
        f"Based on the data '{df['Title_tags'].iloc[0]}', if it contains 'Only one' (even if followed by numbers), there is no issue. "
        f"If the data contains anything other than 'Only one', explain how this could negatively affect SEO and the need for improvement."
    )


    title_tags_impact = get_chatgpt_response(title_tags_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{title_tags_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Multiple Instances: Issue and Impact

    multiple_instances_prompt_issue = (
        f"Summarize in a single line how some pages have multiple title tags, which can cause SEO issues. "
        f"Rephrase creatively and keep it concise. "
        f"Based on the data '{df['Title_tags'].iloc[0]}', if it contains 'Only one' (even if followed by numbers), there is no issue with multiple title tags. "
        f"If the data contains anything other than 'Only one', mention that multiple title tags exist and could cause SEO problems. if the data is empty or zero or null then say that no issues till now!"
    )

    multiple_instances_issue = get_chatgpt_response(multiple_instances_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Multiple Instances: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {multiple_instances_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    multiple_instances_impact_prompt = (
        f"Summarize in a single line the impact of having multiple title tags on a page. "
        f"Impact: [Explain the impact, e.g., 'Multiple title tags on a single page can cause confusion for search engines, diluting the relevance of the targeted keywords and potentially leading to lower rankings.'] using different wordings. "
        f"Based on the data '{df['Title_tags'].iloc[0]}', if it contains 'Only one' (even if followed by numbers), there is no issue with multiple title tags. "
        f"If the data contains anything other than 'Only one', explain the negative impact of having multiple title tags on SEO, such as confusion for search engines and dilution of keyword relevance."
    )

    multiple_instances_impact = get_chatgpt_response(multiple_instances_impact_prompt)
    Impact_para = doc.add_paragraph(style = 'List Bullet')
    Impact_para.add_run('Impact: ').bold = True
    Impact_para.add_run(f"{multiple_instances_impact}")
    Impact_para.paragraph_format.left_indent = Inches(1)
    Impact_para.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, 'Meta Descriptions', level=3)

    # Meta Descriptions: Issue and Impact
    add_bullet_bold(doc, 'Meta Descriptions: ', level=2)

    meta_descriptions_prompt_issue = (
        f"Summarize in a single line how some pages have meta descriptions that are too short, potentially impacting SEO. "
        f"Include the provided meta description data: {df['Meta Description Issues'].iloc[0]}. "
        f"Rephrase creatively. Please note that if the data contains only 'Only one' (even if followed by numbers) just mention that there is no issue with the Meta descriptions. and if and onnly if the data contains any value indicating more than one, mention that the meta descriptions need improvement. if the data is empty or zero or null then say that no issues till now!"
    )

    meta_descriptions_issue = get_chatgpt_response(meta_descriptions_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {meta_descriptions_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    meta_descriptions_impact_prompt = (
        f"Summarize in a single line the impact of short meta descriptions, which may fail to provide enough information, potentially lowering click-through rates (CTR) in search results. "
        f"Include the provided data: {df['Meta Description Issues'].iloc[0]}. "
        f"Rephrase creatively. If the data contains 'Only one' (even if followed by numbers), there is no issue with the meta descriptions. "
        f"If the data contains 'More than one' or any value indicating more than one, explain how short meta descriptions could negatively impact CTR and SEO."
    )

    meta_descriptions_impact = get_chatgpt_response(meta_descriptions_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {meta_descriptions_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    multiple_meta_descriptions_prompt_issue = (
        f"Summarize in a single line how some pages have multiple meta descriptions, which can cause SEO issues. "
        f"Include the provided meta description data: {df['Meta Description Issues'].iloc[0]}. "
        f"Rephrase creatively. If the data contains 'Only one' (even if followed by numbers), there is no issue with multiple meta descriptions. "
        f"If the data contains 'More than one' or any value indicating more than one, mention that having multiple meta descriptions could cause SEO problems."
    )


    multiple_meta_descriptions_issue = get_chatgpt_response(multiple_meta_descriptions_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Multiple Instances: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {multiple_meta_descriptions_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    multiple_meta_descriptions_impact_prompt = (
        f"Impact: Summarize in a single line how multiple meta descriptions can confuse search engines, potentially leading to suboptimal snippet selection in search results. "
        f"Include the provided data on meta description issues: {df['Meta Description Issues'].iloc[0]}. "
        f"Rephrase creatively. If the data contains 'Only one' (even if followed by numbers), there is no issue with multiple meta descriptions. "
        f"If the data contains 'More than one' or any value indicating more than one, explain how multiple meta descriptions could negatively impact SEO by confusing search engines and affecting snippet selection."
    )

    multiple_meta_descriptions_impact = get_chatgpt_response(multiple_meta_descriptions_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {multiple_meta_descriptions_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'


    # Add heading
    add_heading(doc, 'Headings (H1 Tags)', level=3)

    # Headings (H1 Tags): Issue and Impact
    add_bullet_bold(doc, 'Headings (H1 Tags): ', level=2)

    # H1 Tags: Issue
    h1_tags_prompt_issue = (
        f"I need you to, Include the provided H1 tag data: {df['H1 Tag Issues'].iloc[0]}. "    
        f"Summarize in a single line how some pages are missing H1 tags, potentially affecting SEO. "
        f"Rephrase creatively. Please note that if the data contains only 'Only one' (even if followed by numbers) just mention that there is no issue with the H1 tags. and if and onnly if the data contains any value indicating more than one, mention that the H1 tags need improvement."
    )

    h1_tags_issue = get_chatgpt_response(h1_tags_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {h1_tags_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # H1 Tags: Impact
    h1_tags_impact_prompt = (
        f"Summarize in a single line how missing H1 tags can weaken content structure, making it harder for search engines to understand the page’s relevance to specific keywords. "
        f"Include the provided data: {df['H1 Tag Issues'].iloc[0]}. "
        f"Rephrase creatively. If the data contains 'Only one' (even if followed by numbers), there is no issue with the H1 tags. "
        f"If the data contains 'More than one' or any value indicating more than one, explain how missing or incorrect H1 tags could negatively affect SEO and keyword relevance."
    )

    h1_tags_impact = get_chatgpt_response(h1_tags_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {h1_tags_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Multiple H1 Tags: Issue
    multiple_h1_tags_prompt_issue = (
        f"Summarize in a single line how some pages have multiple H1 tags, which can cause SEO issues. "
        f"Include the provided H1 tag data: {df['H1 Tag Issues'].iloc[0]}. "
        f"Rephrase creatively. If the data contains 'Only one' (even if followed by numbers), there is no issue with multiple H1 tags. "
        f"If the data contains 'More than one' or any value indicating more than one, mention that having multiple H1 tags could cause SEO problems."
    )

    multiple_h1_tags_issue = get_chatgpt_response(multiple_h1_tags_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Multiple Instances: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {multiple_h1_tags_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Multiple H1 Tags: Impact
    multiple_h1_tags_impact_prompt = (
        f"Summarize in a single line how multiple H1 tags can dilute the importance of the main heading, making it difficult for search engines to determine the primary focus of the page. "
        f"Include the provided data: {df['H1 Tag Issues'].iloc[0]}. "
        f"Rephrase creatively. If the data contains 'Only one' (even if followed by numbers), there is no issue with multiple H1 tags. "
        f"If the data contains 'More than one' or any value indicating more than one, explain how multiple H1 tags could negatively affect SEO by confusing search engines and diluting the focus of the page."
    )

    multiple_h1_tags_impact = get_chatgpt_response(multiple_h1_tags_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {multiple_h1_tags_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Add heading
    add_heading(doc, 'URL Structure', level=3)

    # URL Structure: Issue and Impact
    add_bullet_bold(doc, 'URL Structure: ', level=2)

    url_structure_data = df[['Internal Links', 'Redirect Chains', 'Sitemap.xml link', 'Indexed Pages']].iloc[0]

    # Issue Analysis Prompt for URL Structure
    url_structure_prompt_issue = (
        f"Analyze whether the site's URL structure follows SEO best practices or has issues. "
        f"Consider factors such as descriptive URLs, URL length, deep URL nesting, or inconsistent formats. "
        f"Also, refer to the following data: Internal Links - {url_structure_data['Internal Links']}, "
        f"Redirect Chains - {url_structure_data['Redirect Chains']}, "
        f"Sitemap.xml link - {url_structure_data['Sitemap.xml link']}, "
        f"Indexed Pages - {url_structure_data['Indexed Pages']}. "
        f"Provide a summary of any issues in a SINGLE line using different wording."
    )

    url_structure_issue = get_chatgpt_response(url_structure_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {url_structure_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    url_structure_impact_prompt = (
        f"Impact: Explain the general impact of URL structure issues on SEO in a single line. "
        f"For example: 'Non-descriptive URLs can negatively affect both user experience and SEO, as they do not provide clear context about the page’s content.' "
        f"Rephrase this impact using different wording, and consider the following data: "
        f"Internal Links - {url_structure_data['Internal Links']}, "
        f"Redirect Chains - {url_structure_data['Redirect Chains']}, "
        f"Sitemap.xml link - {url_structure_data['Sitemap.xml link']}, "
        f"Indexed Pages - {url_structure_data['Indexed Pages']}. "
        f"Summarize the SEO impact in one line, using these insights."
    )

    url_structure_impact = get_chatgpt_response(url_structure_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {url_structure_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Add heading
    add_heading(doc, 'Internal Linking', level=3)

    # Internal Linking: Issue and Impact
    add_bullet_bold(doc, 'Internal Linking: ', level=2)

    # Internal Linking: Issue
    internal_linking_prompt_issue = (
        f"Summarize in a single line how some pages have a weak internal linking structure, potentially affecting SEO. Include the provided internal linking data: {df['Internal Links'].iloc[0]}, and rephrase creatively."
    )
    internal_linking_issue = get_chatgpt_response(internal_linking_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {internal_linking_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Internal Linking: Impact
    internallink_impact_prompt = (
        f'Impact: Summarize in a single line Explain the impact, e.g., "Weak internal linking can lead to poor distribution of link equity, making it difficult for important pages to rank well in search engines.". Reword this creatively, using the data: {df["Internal Links"].iloc[0]}, this is the number.'
    )

    internal_linking_impact = get_chatgpt_response(internallink_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {internal_linking_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph



    # Add heading
    add_heading(doc, 'Image Optimization', level=3)

    # Image Optimization: Issue and Impact
    add_bullet_bold(doc, 'Image Optimization: ', level=2)

    image_optimization_prompt_issue = (
        f"Summarize in a single line how images may be missing ALT text, potentially affecting SEO. Reword creatively, using the data: The number of images missing alt text: {df['Image Optimization'].iloc[0]}, Also include the data!"
    )
    image_optimization_issue = get_chatgpt_response(image_optimization_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {image_optimization_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    image_optimization_impact_prompt = (
        f'Impact: Summarize in a single line the effect of missing ALT text on SEO, such as reduced image search visibility and accessibility, potentially affecting traffic. Reword this creatively, using the data: {df["Image Optimization"].iloc[0]}'
    )

    image_optimization_impact = get_chatgpt_response(image_optimization_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {image_optimization_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, '3. Backlink Analysis', level=3)

    # Domain Strength: Issue and Impact
    add_bullet_bold(doc, 'Domain Strength: ', level=2)

    # Domain Strength: Issue
    domain_strength_prompt_issue = (
        f"Summarize in a single line how the domain authority compares to competitors (e.g., 'The domain authority is average compared to stronger competitors'), including the provided data: {df['Domain Strength/Authority'].iloc[0]}, and rephrase creatively."
    )
    domain_strength_issue = get_chatgpt_response(domain_strength_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {domain_strength_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Domain Strength: Impact
    domain_strength_impact_prompt = (
        f"Impact: Summarize in a single line how average domain authority can make it harder to compete for high-traffic keywords, potentially limiting overall SEO performance, using the provided data: {df['Domain Strength/Authority'].iloc[0]}, and rephrase creatively."
    )
    domain_strength_impact = get_chatgpt_response(domain_strength_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {domain_strength_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Page Strength: Issue and Impact
    add_bullet_bold(doc, 'Page Strength: ', level=2)

    # Page Strength: Issue
    page_strength_prompt_issue = (
        f"Summarize in a single line how the page strength may be low (e.g., 'Pages have low page strength'), including the provided data: {df['Page Strength'].iloc[0]}, and rephrase creatively."
    )
    page_strength_issue = get_chatgpt_response(page_strength_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {page_strength_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Page Strength: Impact
    page_strength_impact_prompt = (
        f"Impact: Summarize in a single line how low page strength can result in poor rankings for important pages, reducing their visibility and conversions, using the provided data: {df['Page Strength'].iloc[0]}, and rephrase creatively."
    )
    page_strength_impact = get_chatgpt_response(page_strength_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {page_strength_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Nofollow vs. Dofollow Links: Issue and Impact
    add_bullet_bold(doc, 'Nofollow vs. Dofollow Links: ', level=2)

    # Nofollow vs. Dofollow Links: Issue
    nofollow_dofollow_prompt_issue = (
        f"Summarize in a single line how the site’s nofollow and dofollow backlinks are distributed (e.g., 'The site has a high proportion of nofollow backlinks'), using the provided data: {df['Nofollow Links'].iloc[0]}, {df['Dofollow Links'].iloc[0]}, and rephrase creatively."
    )
    nofollow_dofollow_issue = get_chatgpt_response(nofollow_dofollow_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {nofollow_dofollow_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Nofollow vs. Dofollow Links: Impact
    nofollow_dofollow_impact_prompt = (
        f"Impact: Summarize in a single line how having more nofollow links might limit SEO performance, as they do not contribute to domain authority, using the provided data: {df['Nofollow Links'].iloc[0]}, {df['Dofollow Links'].iloc[0]}, and rephrase creatively."
    )
    nofollow_dofollow_impact = get_chatgpt_response(nofollow_dofollow_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {nofollow_dofollow_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Add heading
    add_heading(doc, '4. Core Web Vitals', level=3)

    # Time to First Byte (TTFB): Issue and Impact
    add_bullet_bold(doc, 'Time to First Byte (TTFB): ', level=2)

    # Time to First Byte (TTFB): Issue
    ttfb_prompt_issue = (
        f"Summarize in a single line how TTFB might be higher than recommended (e.g., 'The homepage has a TTFB that is slower than optimal'), using the provided data in seconds: {df['Time to first byte distribution'].iloc[0]}, and rephrase creatively."
    )
    ttfb_issue = get_chatgpt_response(ttfb_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {ttfb_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Time to First Byte (TTFB): Impact
    ttfb_impact_prompt = (
        f"Impact: Summarize in a single line how high TTFB can hurt page load speed, leading to worse user experience and potentially harming search rankings, using the provided data, in seconds {df['Time to first byte distribution'].iloc[0]}, and rephrase creatively."
    )
    ttfb_impact = get_chatgpt_response(ttfb_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {ttfb_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Page Load Times: Issue and Impact
    add_bullet_bold(doc, 'Page Load Times: ', level=2)

    # Page Load Times: Issue
    page_load_times_prompt_issue = (
        f"Summarize in a single line how slow page load times are affecting the site (e.g., 'Pages are loading slowly, which could lead to poor user experience'), using the provided data, in seconds {df['Load Time distribution'].iloc[0]}, and rephrase creatively."
    )
    page_load_times_issue = get_chatgpt_response(page_load_times_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {page_load_times_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Page Load Times: Impact
    page_load_times_impact_prompt = (
        f"Impact: Summarize in a single line how slow load times can result in higher bounce rates, lower engagement, and poorer SEO rankings, using the provided data, in seconds  {df['Load Time distribution'].iloc[0]}, and rephrase creatively."
    )
    page_load_times_impact = get_chatgpt_response(page_load_times_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {page_load_times_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # File Size Distribution: Issue and Impact
    add_bullet_bold(doc, 'File Size Distribution: ', level=2)

    # File Size Distribution: Issue
    file_size_distribution_prompt_issue = (
        f"Summarize in a single line how file sizes could be larger than needed (e.g., 'Pages have unnecessarily large file sizes that could impact performance'), using the provided data, in seconds {df['File Size Distribution'].iloc[0]}, and rephrase creatively."
    )
    file_size_distribution_issue = get_chatgpt_response(file_size_distribution_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {file_size_distribution_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # File Size Distribution: Impact
    file_size_distribution_impact_prompt = (
        f"Impact: Summarize in a single line how large file sizes can slow down page load times, negatively impacting user experience and SEO, using the provided data, in seconds {df['File Size Distribution'].iloc[0]}, and rephrase creatively."
    )
    file_size_distribution_impact = get_chatgpt_response(file_size_distribution_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {file_size_distribution_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    # Add heading
    add_heading(doc, '5. Website Usability', level=3)

    # Overall Usability Grade: Issue and Impact
    add_bullet_bold(doc, f'Overall Usability Grade: {df['UsabilityGrade'].iloc[0]}', level=2)


    usability_grade_impact_prompt = (
        'Impact: [Explain the impact, e.g., "A moderate usability grade can indicate areas where user experience may be lacking, leading to higher bounce rates and lower conversion rates."], in a single line, using different wordings'
    )
    usability_grade_impact = get_chatgpt_response(usability_grade_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {usability_grade_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # Device Rendering: Issue and Impact
    add_bullet_bold(doc, 'Device Rendering: ', level=2)

    device_rendering_prompt_issue = (
        f"Summarize in a SINGLE line, in brief,  how the website’s mobile rendering might be suboptimal, focusing on smooth user experience, using the provided data: "
        f"Device Rendering: {df['Device Rendering'].iloc[0]}, Mobile Friendliness Score: {df['Mobile Friendliness Score'].iloc[0]}, "
        f"Cumulative Layout Shift: {df['Cumulatice Layout Shift'].iloc[0]}, and Page Speed Index (mobile): {df['SpeedIndexMobile'].iloc[0]}. "
        f"Rephrase creatively for variety. Mention the Data also"
    )


    device_rendering_issue = get_chatgpt_response(device_rendering_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {device_rendering_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    device_rendering_impact_prompt = (
        'Impact: [Explain the impact, e.g., "Poor mobile rendering can lead to a subpar user experience, higher bounce rates, and lower mobile search rankings."], in a single line, using different wordings'
    )
    device_rendering_impact = get_chatgpt_response(device_rendering_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet', )
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {device_rendering_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    # User Experience: Issue and Impact
    add_bullet_bold(doc, 'User Experience: ', level=2)

    user_experience_prompt_issue = (
        f"Summarize the user experience in a single line based on the provided data. Ensure the data is mentioned clearly. If there are no issues, state explicitly that there are no issues. "
        f"The data is as follows: First Contentful Paint: {df['First Contentful Paint'].iloc[0]}, "
        f"Largest Contentful Paint: {df['Largest Contentful Paint'].iloc[0]}, Total Blocking Time: {df['Total Blocking Time'].iloc[0]}, "
        f"Cumulative Layout Shift: {df['Cumulatice Layout Shift'].iloc[0]}. Mention the data also!"
    )

    user_experience_issue = get_chatgpt_response(user_experience_prompt_issue)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Issue: ')  # Add the bullet title
    bullet_paragraph.add_run(f" {user_experience_issue}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph

    user_experience_impact_prompt = (
        'Impact: [Explain the impact, e.g., "A confusing navigation structure can frustrate users, leading to higher bounce rates and reduced engagement, negatively affecting SEO performance."], in a single line, using different wordings'
    )
    user_experience_impact = get_chatgpt_response(user_experience_impact_prompt)
    bullet_paragraph = doc.add_paragraph(style='List Bullet')
    bullet_paragraph.add_run('Impact: ').bold = True  # Add the bullet title
    bullet_paragraph.add_run(f" {user_experience_impact}")  # Append the response text
    bullet_paragraph.paragraph_format.left_indent = Inches(1)  # Set indentation
    bullet_paragraph.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    heading20=doc.add_heading('6. Conclusion and Key Takeaways', level=3)
    heading20.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run20 = heading20.runs[0]  # Get the run from the heading paragraph
    heading_run20.font.italic = False  # Set font italic to False to remove italics
    heading_run20.font.color.rgb = RGBColor(0,0,0)

    Impact_score = doc.add_paragraph(style='List Bullet 2')
    Impact_score.add_run('Summary of Findings: ').bold = True
    Impact_score.add_run('[Summarize the key issues identified in the audit.]')
    Impact_score.paragraph_format.left_indent = Inches(1)
    Impact_score.style.font.name = 'Calibri'  # Set Calibri font for the entire paragraph


    Impact_score = doc.add_paragraph(style='List Bullet 2')
    Impact_score.add_run('Key Takeaways: ').bold = True
    Impact_score.add_run('[Highlight the importance of addressing these issues for improving SEO performance and overall website health.]')
    Impact_score.paragraph_format.left_indent = Inches(1)
    Impact_score.paragraph_format.name = 'Calibri'

    # Access the first section of the document
    section = doc.sections[0]
    header = section.header

    # Add a new paragraph to the header and center-align it
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the image to the paragraph
    run = paragraph.add_run()
    run.add_picture("Sinopia_logo.png", width=Inches(2))  # Adjust the image width as needed

    # Save the updated document
    the_doc = f'{company_name}.docx'
    doc.save(the_doc)
  
    return the_doc
